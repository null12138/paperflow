#!/usr/bin/env python3
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
ISSUES = json.loads((HERE / 'issues.json').read_text(encoding='utf-8'))
REVIEWS = json.loads((HERE / 'review_results.json').read_text(encoding='utf-8'))
OUT = HERE / '双师课课后作业_逐题纠错高亮对比_内置渲染版.docx'

FONT = 'PingFang SC'
INK = '17324D'
BLUE = '245B78'
MUTED = '667085'
LINE = 'D5DDE5'
RED = 'B42318'
RED_BG = 'FFF3F2'
GREEN = '067647'
GREEN_BG = 'ECFDF3'
PALE_BLUE = 'EAF2F7'
PALE_GRAY = 'F7F9FB'


def set_font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def cell_margins(cell, top=130, start=170, bottom=130, end=170):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar'); tc_pr.append(tc_mar)
    for tag, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tc_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}'); tc_mar.append(node)
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'),'dxa')


def borders(table, color=LINE, size='6'):
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.first_child_found_in('w:tblBorders')
    if old is not None: tbl_pr.remove(old)
    b = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        x=OxmlElement(f'w:{edge}'); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),size)
        x.set(qn('w:space'),'0'); x.set(qn('w:color'),color); b.append(x)
    tbl_pr.append(b)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout=OxmlElement('w:tblLayout'); tbl_pr.append(layout)
    layout.set(qn('w:type'),'fixed')
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    tbl_w.set(qn('w:w'), str(sum(widths_dxa))); tbl_w.set(qn('w:type'),'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind=OxmlElement('w:tblInd'); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'),str(indent)); tbl_ind.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_dxa:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(width)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcw=cell._tc.get_or_add_tcPr().first_child_found_in('w:tcW')
            tcw.set(qn('w:w'),str(widths_dxa[i])); tcw.set(qn('w:type'),'dxa')


def cant_split(row):
    tr_pr=row._tr.get_or_add_trPr(); x=OxmlElement('w:cantSplit'); tr_pr.append(x)


def repeat_header(row):
    tr_pr=row._tr.get_or_add_trPr(); x=OxmlElement('w:tblHeader'); x.set(qn('w:val'),'true'); tr_pr.append(x)


def keep_with_next(p, value=True):
    p.paragraph_format.keep_with_next=value


def add_bottom_border(p, color=BLUE, size='14', space='4'):
    p_pr=p._p.get_or_add_pPr(); pbdr=p_pr.find(qn('w:pBdr'))
    if pbdr is None: pbdr=OxmlElement('w:pBdr'); p_pr.append(pbdr)
    bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),size)
    bottom.set(qn('w:space'),space); bottom.set(qn('w:color'),color); pbdr.append(bottom)


def add_page_field(paragraph):
    run=paragraph.add_run(); fld_begin=OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    fld_sep=OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'),'separate')
    txt=OxmlElement('w:t'); txt.text='1'
    fld_end=OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'),'end')
    run._r.extend([fld_begin,instr,fld_sep,txt,fld_end]); set_font(run,8.5,MUTED)


def style_document(doc):
    sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=Inches(0.78); sec.bottom_margin=Inches(0.72)
    sec.left_margin=Inches(1); sec.right_margin=Inches(1)
    sec.header_distance=Inches(0.38); sec.footer_distance=Inches(0.38)

    normal=doc.styles['Normal']; normal.font.name=FONT; normal._element.rPr.rFonts.set(qn('w:eastAsia'),FONT)
    normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,before,after,color in [('Title',27,0,8,INK),('Heading 1',17,18,9,BLUE),('Heading 2',12.5,10,5,INK)]:
        s=doc.styles[name]; s.font.name=FONT; s._element.rPr.rFonts.set(qn('w:eastAsia'),FONT)
        s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
        s.paragraph_format.keep_with_next=True

    hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    set_font(hp.add_run('双师课课后作业  /  逐题纠错'),8.5,MUTED)
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_font(fp.add_run('逐题审阅 · 修改前后高亮对比   |   '),8.5,MUTED); add_page_field(fp)


def add_stat_strip(doc):
    table=doc.add_table(rows=1,cols=3); table.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table,[3120,3120,3120],indent=120); borders(table,PALE_BLUE,'4')
    values=[('17','章节'),('67','题逐题审阅'),('30','处修改')]
    for cell,(num,label) in zip(table.rows[0].cells,values):
        shade(cell,PALE_BLUE); cell_margins(cell,150,140,150,140); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
        set_font(p.add_run(num),18,BLUE,True); set_font(p.add_run(f'\n{label}'),9,MUTED,False)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def add_compare_card(doc, idx, issue):
    p=doc.add_paragraph(style='Heading 2'); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    set_font(p.add_run(f'{idx:02d}  第{issue["section"]}节 · 第{issue["question"]}题'),12.5,INK,True)
    set_font(p.add_run(f'   原稿约第{issue["page"]}页'),9,MUTED)
    keep_with_next(p)

    table=doc.add_table(rows=2,cols=2); table.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table,[4680,4680],indent=120); borders(table,LINE,'6')
    labels=[('修改前',RED,RED_BG),('修改后',GREEN,GREEN_BG)]
    for i,(label,color,fill) in enumerate(labels):
        c=table.cell(0,i); shade(c,fill); cell_margins(c,95,170,85,170)
        pp=c.paragraphs[0]; pp.paragraph_format.space_after=Pt(0); set_font(pp.add_run(label),9,color,True)
        c2=table.cell(1,i); shade(c2,fill); cell_margins(c2,165,170,165,170); c2.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        pp=c2.paragraphs[0]; pp.paragraph_format.space_after=Pt(0); pp.paragraph_format.line_spacing=1.22
        set_font(pp.add_run(issue['before'] if i==0 else issue['after']),10.5,color,False)
    for row in table.rows: cant_split(row)

    note=doc.add_paragraph(); note.paragraph_format.space_before=Pt(4); note.paragraph_format.space_after=Pt(7)
    note.paragraph_format.left_indent=Inches(0.08); note.paragraph_format.right_indent=Inches(0.08)
    set_font(note.add_run('问题  '),8.8,BLUE,True); set_font(note.add_run(issue['problem']),8.8,MUTED)
    set_font(note.add_run('   校核  '),8.8,BLUE,True); set_font(note.add_run(issue['basis']),8.8,MUTED)


def add_coverage(doc):
    doc.add_page_break()
    p=doc.add_paragraph(style='Heading 1'); p.paragraph_format.space_before=Pt(0)
    set_font(p.add_run('67题逐题审阅覆盖表'),19,INK,True); add_bottom_border(p)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10)
    set_font(p.add_run('“通过”表示题干、选项、答案和解析均已逐题核对，未发现需要修改之处；“需修改”对应前文编号。'),9.5,MUTED)
    table=doc.add_table(rows=1,cols=5); table.alignment=WD_TABLE_ALIGNMENT.LEFT
    widths=[700,2500,900,1200,4060]; set_table_geometry(table,widths,indent=120); borders(table,LINE,'5')
    headers=['顺序','定位','原稿页','结论','对应修改项']
    for cell,text in zip(table.rows[0].cells,headers):
        shade(cell,BLUE); cell_margins(cell,70,100,70,100); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
        set_font(p.add_run(text),8,'FFFFFF',True)
    repeat_header(table.rows[0]); cant_split(table.rows[0])
    for x in REVIEWS:
        row=table.add_row(); cant_split(row)
        refs='、'.join(f'{i:02d}' for i in x['issues']) if x['issues'] else '—'
        vals=[str(x['seq']),f'第{x["section"]}节 第{x["question"]}题',str(x['page']),x['status'],refs]
        fill=RED_BG if x['status']=='需修改' else ('FFFFFF' if x['seq']%2 else PALE_GRAY)
        for i,(cell,val) in enumerate(zip(row.cells,vals)):
            shade(cell,fill); cell_margins(cell,32,90,32,90); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
            set_font(p.add_run(val),7.2,RED if (x['status']=='需修改' and i>=3) else INK, x['status']=='需修改' and i==3)


def main():
    doc=Document(); style_document(doc)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(26); p.paragraph_format.space_after=Pt(7)
    set_font(p.add_run('逐题审阅报告'),10,BLUE,True)
    p=doc.add_paragraph(style='Title'); p.paragraph_format.space_after=Pt(5)
    set_font(p.add_run('双师课课后作业'),27,INK,True)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18)
    set_font(p.add_run('修改前后高亮对比'),15,BLUE,True)
    add_stat_strip(doc)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14)
    set_font(p.add_run('阅读说明  '),9.5,INK,True)
    set_font(p.add_run('仅列出发现问题的题。左侧浅红为原文，右侧浅绿为修正版；末尾附全部67题覆盖表。'),9.5,MUTED)
    add_bottom_border(p,LINE,'6','7')

    current=None
    for idx,issue in enumerate(ISSUES,1):
        if issue['section']!=current:
            current=issue['section']
            p=doc.add_paragraph(style='Heading 1'); set_font(p.add_run(f'第{current}节'),17,BLUE,True); add_bottom_border(p,PALE_BLUE,'9','5')
        add_compare_card(doc,idx,issue)
    add_coverage(doc)
    props=doc.core_properties; props.title='双师课课后作业：逐题纠错高亮对比'; props.subject='17节67题逐题审阅'; props.author='逐题审阅'
    doc.save(OUT); print(OUT)


if __name__=='__main__': main()
